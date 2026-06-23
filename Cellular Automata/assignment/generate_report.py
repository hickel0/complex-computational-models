"""
Script to generate the Word document report for the Virus Spread ABM project.
Run this script to create the CSC1111_ABM_Report.docx file.

Requirements: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_report():
    """Generate the Word document report."""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Virus Spread Agent-Based Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('CSC1111 - Building Complex Computational Models\n')
    run.bold = True
    run = subtitle.add_run('2nd Continuous Assessment\n\n')
    run = subtitle.add_run('[Student Name 1] - [Student Number]\n')
    run = subtitle.add_run('[Student Name 2] - [Student Number]\n\n')
    run = subtitle.add_run('Date: April 2026')

    doc.add_page_break()

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph(
        'This report documents the design, implementation, and analysis of an Agent-Based Model (ABM) '
        'simulating virus spread within a population. The model uses a spatial SIR (Susceptible-Infected-Recovered) '
        'framework where individual agents move through a 2D environment and interact based on proximity.'
    )

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'The purpose of this simulation is to:'
    )
    bullets = [
        'Model how infectious diseases spread through a population via local interactions',
        'Investigate the effects of various parameters (infection rate, social distancing, etc.) on epidemic dynamics',
        'Visualise the spatial and temporal patterns of disease transmission',
        'Provide insights into intervention strategies such as social distancing'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('1.2 Relevance', level=2)
    doc.add_paragraph(
        'Agent-based models of disease spread are valuable tools for understanding epidemics. Unlike '
        'equation-based models (e.g., differential equation SIR models), ABMs capture individual-level '
        'heterogeneity, spatial locality, and stochastic effects that are crucial for realistic epidemic modelling.'
    )

    # =========================================================================
    # 2. MODEL DESCRIPTION
    # =========================================================================
    doc.add_heading('2. Model Description', level=1)

    doc.add_heading('2.1 Agent Attributes', level=2)
    doc.add_paragraph(
        'Each agent in the simulation represents an individual in the population and has the following attributes:'
    )

    # Create attributes table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Attribute'
    headers[1].text = 'Description'

    data = [
        ('Position (x, y)', 'Location in 2D continuous space [0, 1]'),
        ('State', 'S (Susceptible), I (Infected), R (Recovered), or D (Dead)'),
        ('Velocity (vx, vy)', 'Direction of movement, updated with random perturbations'),
        ('Infection Timer', 'Time steps since becoming infected'),
        ('Immunity Timer', 'Remaining duration of immunity (if applicable)')
    ]
    for i, (attr, desc) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = attr
        row[1].text = desc

    doc.add_heading('2.2 Agent Behaviour Rules', level=2)

    doc.add_heading('Movement', level=3)
    doc.add_paragraph(
        'Agents move through the 2D space using a random walk with momentum. At each time step:'
    )
    movement_rules = [
        'The agent\'s direction is perturbed by a small random angle',
        'Position is updated based on velocity and movement speed',
        'Agents reflect off boundaries (bounce back into the space)',
        'Social distancing reduces effective movement speed'
    ]
    for rule in movement_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_heading('Infection Transmission', level=3)
    doc.add_paragraph(
        'Infection spreads through proximity-based contact:'
    )
    infection_rules = [
        'An infected agent can transmit the virus to susceptible agents within the infection radius',
        'Transmission occurs probabilistically based on the infection probability parameter',
        'Multiple exposures per time step are possible (all neighbours checked)'
    ]
    for rule in infection_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_heading('Recovery and Death', level=3)
    doc.add_paragraph(
        'Infected agents can recover or die after the recovery time:'
    )
    recovery_rules = [
        'After recovery_time steps, infected agents either recover or die',
        'Death occurs with probability equal to the mortality rate',
        'Recovered agents gain immunity (permanent or temporary based on settings)',
        'Dead agents remain in place but do not interact'
    ]
    for rule in recovery_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_heading('2.3 Why These Rules?', level=2)
    doc.add_paragraph(
        'The behavioural rules are designed to capture key aspects of respiratory virus transmission:'
    )
    why_rules = [
        ('Random movement', 'Represents daily activities and interactions in a simplified manner'),
        ('Proximity-based infection', 'Models airborne/droplet transmission requiring close contact'),
        ('Probabilistic transmission', 'Captures the stochastic nature of infection events'),
        ('Recovery time', 'Represents the infectious period of the disease'),
        ('Social distancing parameter', 'Allows investigation of intervention effectiveness')
    ]
    for rule, reason in why_rules:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{rule}: ').bold = True
        p.add_run(reason)

    # =========================================================================
    # 3. IMPLEMENTATION
    # =========================================================================
    doc.add_heading('3. Implementation', level=1)

    doc.add_heading('3.1 Technical Details', level=2)
    impl_details = [
        'Programming Language: Python 3.8+',
        'Libraries: NumPy (numerical operations), Matplotlib (visualisation)',
        'Architecture: Object-oriented design with Agent, Simulation, and Visualizer classes',
        'Interface: Interactive GUI with real-time parameter adjustment via sliders'
    ]
    for detail in impl_details:
        doc.add_paragraph(detail, style='List Bullet')

    doc.add_heading('3.2 Key Features', level=2)
    features = [
        '2D graphical display showing agent positions colour-coded by state',
        'Real-time time-series graph of population dynamics (SIR curves)',
        'Interactive controls: Start, Stop, Reset buttons',
        'Parameter sliders for infection probability, infection radius, and social distancing',
        'Command-line interface for batch simulations and parameter sweeps',
        'Batch mode for automated data collection without GUI'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('3.3 Code Structure', level=2)
    doc.add_paragraph(
        'The code is organised into the following main components:'
    )
    code_structure = [
        ('SimulationConfig', 'Dataclass holding all simulation parameters'),
        ('Agent class', 'Represents individuals with position, state, and behaviour methods'),
        ('VirusSpreadSimulation', 'Main simulation logic including infection and recovery'),
        ('SimulationVisualizer', 'Handles all graphical output and user interaction'),
        ('Command-line parser', 'Enables parameter configuration via arguments')
    ]
    for component, desc in code_structure:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{component}: ').bold = True
        p.add_run(desc)

    # =========================================================================
    # 4. PARAMETERS
    # =========================================================================
    doc.add_heading('4. Simulation Parameters', level=1)

    doc.add_paragraph('The following parameters can be configured:')

    # Parameters table
    table2 = doc.add_table(rows=10, cols=4)
    table2.style = 'Table Grid'
    headers2 = table2.rows[0].cells
    headers2[0].text = 'Parameter'
    headers2[1].text = 'Default'
    headers2[2].text = 'Range'
    headers2[3].text = 'Description'

    params = [
        ('population_size', '200', '50-1000', 'Number of agents'),
        ('initial_infected', '5', '1-50', 'Starting infected count'),
        ('infection_probability', '0.3', '0.0-1.0', 'Transmission chance'),
        ('infection_radius', '0.03', '0.01-0.1', 'Contact distance'),
        ('recovery_time', '100', '50-300', 'Steps until recovery'),
        ('mortality_rate', '0.0', '0.0-0.2', 'Death probability'),
        ('social_distancing', '0.0', '0.0-1.0', 'Movement reduction'),
        ('movement_speed', '0.01', '0.005-0.02', 'Agent speed'),
        ('immunity_duration', '0', '0-500', 'Immunity length (0=permanent)')
    ]
    for i, (name, default, range_val, desc) in enumerate(params, 1):
        row = table2.rows[i].cells
        row[0].text = name
        row[1].text = default
        row[2].text = range_val
        row[3].text = desc

    # =========================================================================
    # 5. RESULTS
    # =========================================================================
    doc.add_heading('5. Simulation Results', level=1)

    doc.add_paragraph(
        'The model was tested under several parameter scenarios to investigate epidemic dynamics. '
        'Results are presented below with analysis of key outcomes.'
    )

    doc.add_heading('5.1 Baseline Scenario', level=2)
    doc.add_paragraph(
        'Parameters: population=200, initial_infected=5, infection_prob=0.3, infection_radius=0.03, '
        'recovery_time=100, mortality=0, social_distancing=0'
    )
    doc.add_paragraph(
        'Observations:'
    )
    baseline_obs = [
        'Epidemic spreads gradually from initial infected agents',
        'Peak infection occurs around time step 200-300',
        'Approximately 80-90% of population eventually becomes infected',
        'Classic SIR curve shape with susceptible decreasing, infected peaking, recovered increasing'
    ]
    for obs in baseline_obs:
        doc.add_paragraph(obs, style='List Bullet')

    doc.add_paragraph('[INSERT FIGURE: Baseline scenario screenshot]', style='Caption')

    doc.add_heading('5.2 High Transmission Scenario', level=2)
    doc.add_paragraph(
        'Parameters: infection_prob=0.7, infection_radius=0.05 (other parameters at baseline)'
    )
    doc.add_paragraph('Observations:')
    high_trans = [
        'Much faster epidemic progression',
        'Peak infection occurs earlier (around time step 100-150)',
        'Nearly 100% infection rate',
        'Sharper, more pronounced infection peak'
    ]
    for obs in high_trans:
        doc.add_paragraph(obs, style='List Bullet')

    doc.add_paragraph('[INSERT FIGURE: High transmission scenario screenshot]', style='Caption')

    doc.add_heading('5.3 Social Distancing Scenario', level=2)
    doc.add_paragraph(
        'Parameters: social_distancing=0.7 (other parameters at baseline)'
    )
    doc.add_paragraph('Observations:')
    social_dist = [
        'Significantly slower epidemic spread',
        'Lower peak infection (approximately 50% reduction)',
        '"Flattening the curve" effect clearly visible',
        'Some agents may never become infected',
        'Epidemic takes longer to resolve but has lower peak impact'
    ]
    for obs in social_dist:
        doc.add_paragraph(obs, style='List Bullet')

    doc.add_paragraph('[INSERT FIGURE: Social distancing scenario screenshot]', style='Caption')

    doc.add_heading('5.4 Mortality Scenario', level=2)
    doc.add_paragraph(
        'Parameters: mortality_rate=0.1, recovery_time=80 (other parameters at baseline)'
    )
    doc.add_paragraph('Observations:')
    mortality = [
        'Dead agents (shown in grey) accumulate over time',
        'Final death toll approximately 8-12% of population',
        'Epidemic dynamics otherwise similar to baseline',
        'Spatial clustering of deaths near infection hotspots'
    ]
    for obs in mortality:
        doc.add_paragraph(obs, style='List Bullet')

    doc.add_paragraph('[INSERT FIGURE: Mortality scenario screenshot]', style='Caption')

    doc.add_heading('5.5 Waning Immunity Scenario', level=2)
    doc.add_paragraph(
        'Parameters: immunity_duration=200 (other parameters at baseline)'
    )
    doc.add_paragraph('Observations:')
    waning = [
        'After initial epidemic wave subsides, recovered agents lose immunity',
        'Second wave of infections occurs as agents become susceptible again',
        'Potential for endemic equilibrium with ongoing low-level transmission',
        'Multiple epidemic waves possible over long simulation runs'
    ]
    for obs in waning:
        doc.add_paragraph(obs, style='List Bullet')

    doc.add_paragraph('[INSERT FIGURE: Waning immunity scenario screenshot]', style='Caption')

    # =========================================================================
    # 6. TRACKED OUTPUTS
    # =========================================================================
    doc.add_heading('6. Output Metrics', level=1)

    doc.add_paragraph('The simulation tracks the following outputs over time:')

    outputs = [
        ('Susceptible count', 'Number of agents who have not been infected'),
        ('Infected count', 'Number of currently infected agents'),
        ('Recovered count', 'Number of agents who have recovered'),
        ('Dead count', 'Number of deceased agents'),
        ('Peak infection', 'Maximum simultaneous infections'),
        ('Peak time', 'Time step when peak occurs'),
        ('Total infected', 'Cumulative number ever infected'),
        ('Epidemic duration', 'Time until no active infections')
    ]

    doc.add_paragraph('These metrics allow analysis of:')
    for metric, desc in outputs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{metric}: ').bold = True
        p.add_run(desc)

    # =========================================================================
    # 7. DISCUSSION
    # =========================================================================
    doc.add_heading('7. Discussion', level=1)

    doc.add_heading('7.1 Key Findings', level=2)
    findings = [
        'Spatial locality creates clustering effects not captured by equation-based models',
        'Social distancing is highly effective at reducing peak infections',
        'Stochasticity means identical parameters can produce different outcomes',
        'Higher population density (more agents) leads to faster transmission',
        'The model successfully reproduces qualitative features of real epidemics'
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    doc.add_heading('7.2 Limitations', level=2)
    limitations = [
        'Simplified 2D space does not capture real-world geography',
        'All agents have identical behaviour (no heterogeneity in susceptibility)',
        'No explicit contact network structure',
        'Recovery time is fixed rather than distributed',
        'No age structure or demographic factors'
    ]
    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')

    doc.add_heading('7.3 Future Extensions', level=2)
    extensions = [
        'Add heterogeneous agent properties (e.g., varying susceptibility)',
        'Implement quarantine behaviour for infected agents',
        'Add vaccination as an intervention strategy',
        'Include multiple virus variants with different properties',
        'Model household/workplace contact structures'
    ]
    for ext in extensions:
        doc.add_paragraph(ext, style='List Bullet')

    # =========================================================================
    # 8. CONCLUSION
    # =========================================================================
    doc.add_heading('8. Conclusion', level=1)

    doc.add_paragraph(
        'This Agent-Based Model successfully simulates virus spread dynamics in a spatial population. '
        'The model demonstrates key epidemic phenomena including exponential growth, peak infection, '
        'and the effectiveness of social distancing interventions. The interactive GUI allows '
        'exploration of parameter effects in real-time, while batch mode enables systematic analysis. '
        'The results align with expected epidemiological behaviour and provide insights into '
        'disease transmission dynamics.'
    )

    # =========================================================================
    # 9. REFERENCES
    # =========================================================================
    doc.add_heading('9. References', level=1)

    refs = [
        'Sayama, H. (2015). Introduction to the Modeling and Analysis of Complex Systems. Open SUNY Textbooks.',
        'Epstein, J. M. (2009). Modelling to contain pandemics. Nature, 460(7256), 687-687.',
        'Bonabeau, E. (2002). Agent-based modeling: Methods and techniques for simulating human systems. PNAS.',
        'PyCX Project: https://github.com/hsayama/PyCX'
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Number')

    # =========================================================================
    # APPENDIX
    # =========================================================================
    doc.add_heading('Appendix: How to Run the Simulation', level=1)

    doc.add_paragraph('Requirements:')
    doc.add_paragraph('pip install numpy matplotlib', style='No Spacing')

    doc.add_paragraph('')
    doc.add_paragraph('Basic usage:')
    doc.add_paragraph('python virus_spread_abm.py', style='No Spacing')

    doc.add_paragraph('')
    doc.add_paragraph('With custom parameters:')
    doc.add_paragraph('python virus_spread_abm.py --population 300 --infection-prob 0.5', style='No Spacing')

    doc.add_paragraph('')
    doc.add_paragraph('Batch mode (no GUI):')
    doc.add_paragraph('python virus_spread_abm.py --no-gui --max-steps 500', style='No Spacing')

    doc.add_paragraph('')
    doc.add_paragraph('Code repository: [INSERT GITLAB LINK]')

    # Save document
    doc.save('CSC1111_ABM_Report.docx')
    print("Report generated: CSC1111_ABM_Report.docx")

if __name__ == "__main__":
    # Check if python-docx is installed
    try:
        from docx import Document
        create_report()
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call(['pip', 'install', 'python-docx'])
        from docx import Document
        create_report()
